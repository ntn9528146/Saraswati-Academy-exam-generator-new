import streamlit as st
import os
import re
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ALL_CLASSES = [
    "Nursery", "LKG", "UKG",
    "Class 1", "Class 2", "Class 3", "Class 4", "Class 5",
    "Class 6", "Class 7", "Class 8", "Class 9", "Class 10",
    "Class 11", "Class 12"
]

PRIMARY_MIDDLE_SUBJECTS = sorted([
    "English", "Hindi", "Mathematics", "Environmental Studies (EVS)", 
    "Science", "Social Science", "Computer Studies", "General Knowledge (GK)", "Sanskrit"
])

SECONDARY_SUBJECTS = sorted([
    "Information Technology (Code 402)", "Artificial Intelligence (Code 417)",
    "Mathematics (Standard - 041)", "Mathematics (Basic - 241)", "Science (086)", 
    "Social Science (087)", "English Language & Literature (184)", 
    "Hindi Course A (002)", "Hindi Course B (085)", "Sanskrit (122)"
])

SENIOR_SECONDARY_SUBJECTS = sorted([
    "Informatics Practices (Code No. 065)", "Computer Science (Code No. 083)",
    "Mathematics (041)", "Applied Mathematics (241)", "Physics (042)", 
    "Chemistry (043)", "Biology (044)", "Accountancy (055)", "Business Studies (054)", 
    "Economics (030)", "English Core (301)", "History (027)", 
    "Political Science (028)", "Geography (029)", "Physical Education (048)"
])

JUNIOR_Q_TYPES = [
    "Multiple Choice Questions (MCQ)", "Fill in the Blanks", "True / False",
    "Match the Following", "Very Short Answer (VSA)", "Short Answer (SA)",
    "Long Answer (LA)", "Picture / Passage Based Questions"
]

DEFAULT_JUNIOR_MARKS = {
    "Multiple Choice Questions (MCQ)": 1, "Fill in the Blanks": 1, "True / False": 1,
    "Match the Following": 2, "Very Short Answer (VSA)": 2, "Short Answer (SA)": 3,
    "Long Answer (LA)": 5, "Picture / Passage Based Questions": 4
}

# ---------------- DOCX FORMATTING HELPERS ----------------
def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders %s>'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tcBorders>' % nsdecls('w')
    )
    tcPr.append(tcBorders)

def add_bold_runs(paragraph, text):
    tokens = re.split(r'(\*\*.*?\*\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        else:
            paragraph.add_run(token)

def export_cbse_sqp_docx(school_name, class_level, subject, total_marks, time_allowed, raw_content, logo_path=None):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)

    if logo_path and os.path.exists(logo_path):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.paragraph_format.space_after = Pt(2)
            p_logo.add_run().add_picture(logo_path, width=Inches(0.85))
        except Exception:
            pass

    p_sch = doc.add_paragraph()
    p_sch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sch.paragraph_format.space_after = Pt(2)
    r1 = p_sch.add_run(school_name.upper())
    r1.bold = True
    r1.font.size = Pt(14)

    p_sess = doc.add_paragraph()
    p_sess.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sess.paragraph_format.space_after = Pt(2)
    r2 = p_sess.add_run(f"SAMPLE QUESTION PAPER (SESSION 2026-2027)\n{class_level.upper()} — {subject.upper()}")
    r2.bold = True
    r2.font.size = Pt(12)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(3.5)
    tbl.columns[1].width = Inches(3.5)

    c1 = tbl.cell(0, 0).paragraphs[0]
    rc1 = c1.add_run(f"Time Allowed: {time_allowed}")
    rc1.bold = True

    c2 = tbl.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rc2 = c2.add_run(f"Maximum Marks: {total_marks}")
    rc2.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    lines = [l.strip() for l in raw_content.split("\n") if l.strip() and l.strip() != "---"]
    in_gi = True
    gi_lines = []
    question_lines = []

    for line in lines:
        if line.startswith("### SECTION") or line.startswith("SECTION A") or line.startswith("**SECTION"):
            in_gi = False
        if in_gi:
            gi_lines.append(line)
        else:
            question_lines.append(line)

    for gl in gi_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if "GENERAL INSTRUCTIONS" in gl.upper():
            r = p.add_run(gl.replace("###", "").replace("**", "").strip())
            r.bold = True
            r.underline = True
        else:
            add_bold_runs(p, gl)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Official 3-Column Table
    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_widths = [Inches(0.8), Inches(5.4), Inches(0.8)]

    curr_q_no = ""
    curr_q_text = []
    curr_marks = ""

    def flush_question(q_no, q_text_list, marks):
        if not q_no and not q_text_list:
            return
        row = table.add_row()
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            set_cell_border(row.cells[idx])

        c0 = row.cells[0].paragraphs[0]
        c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c0.add_run(q_no)
        r.bold = True

        c1 = row.cells[1].paragraphs[0]
        c1.paragraph_format.space_before = Pt(1)
        c1.paragraph_format.space_after = Pt(1)
        for idx, t in enumerate(q_text_list):
            if idx > 0:
                p = row.cells[1].add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                add_bold_runs(p, t)
            else:
                add_bold_runs(c1, t)

        c2 = row.cells[2].paragraphs[0]
        c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_m = c2.add_run(marks)
        r_m.bold = True

    for ql in question_lines:
        if "SECTION" in ql.upper() and (ql.startswith("###") or ql.startswith("**") or len(ql) < 70):
            flush_question(curr_q_no, curr_q_text, curr_marks)
            curr_q_no, curr_q_text, curr_marks = "", [], ""
            
            sec_row = table.add_row()
            cell_merged = sec_row.cells[0]
            cell_merged.merge(sec_row.cells[1]).merge(sec_row.cells[2])
            set_cell_border(cell_merged)
            
            p_sec = cell_merged.paragraphs[0]
            p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sec.paragraph_format.space_before = Pt(4)
            p_sec.paragraph_format.space_after = Pt(4)
            r = p_sec.add_run(ql.replace("###", "").replace("**", "").strip())
            r.bold = True
            continue

        m = re.match(r"^(Q\s*\.?\s*\d+|\d+)\.?\s*(.*)", ql)
        if m and not ql.startswith("(") and not ql.startswith("-"):
            flush_question(curr_q_no, curr_q_text, curr_marks)
            curr_q_no = m.group(1).replace("Q", "").replace(".", "").strip() + "."
            rest_text = m.group(2).strip()
            m_marks = re.search(r"\[(\d+)\s*(?:Marks|Mark)?\]$", rest_text)
            if m_marks:
                curr_marks = m_marks.group(1)
                rest_text = rest_text[:m_marks.start()].strip()
            else:
                curr_marks = "1"
            curr_q_text = [rest_text] if rest_text else []
        else:
            if curr_q_no:
                curr_q_text.append(ql)

    flush_question(curr_q_no, curr_q_text, curr_marks)

    os.makedirs("generated_papers", exist_ok=True)
    clean_sub = re.sub(r'[^a-zA-Z0-9_]', '_', subject)
    clean_cls = re.sub(r'[^a-zA-Z0-9_]', '_', class_level)
    fname = f"{clean_cls}_{clean_sub}_Official_CBSE_Paper.docx"
    fpath = os.path.join("generated_papers", fname)
    doc.save(fpath)
    return fpath, fname

# ---------------- BULLETPROOF DYNAMIC AI GENERATOR ----------------
def run_gemini_paper_generator(api_key_str, system_prompt, user_prompt):
    genai.configure(api_key=api_key_str)
    
    # Live discovery of models available on your key
    available_models = []
    try:
        for m in genai.list_models():
            if "generateContent" in m.supported_generation_methods:
                available_models.append(m.name)
    except Exception:
        pass

    fallback_candidates = [
        "models/gemini-2.0-flash", "models/gemini-1.5-flash", "models/gemini-1.5-pro",
        "gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"
    ]
    
    model_pool = []
    for m in available_models:
        if "flash" in m.lower() and m not in model_pool:
            model_pool.append(m)
    for m in available_models + fallback_candidates:
        if m not in model_pool:
            model_pool.append(m)

    last_err = None
    for model_name in model_pool:
        try:
            model = genai.GenerativeModel(
                model_name=model_name,
                system_instruction=system_prompt,
                generation_config={"temperature": 0.2, "max_output_tokens": 8192}
            )
            response = model.generate_content(user_prompt)
            if response and response.text:
                clean_res = response.text.strip()
                clean_res = re.sub(r"^\*Let's .*?\n\n", "", clean_res, flags=re.DOTALL)
                clean_res = re.sub(r"^<think>.*?</think>\n?", "", clean_res, flags=re.DOTALL)
                return clean_res.strip()
        except Exception as err:
            last_err = err
            continue
            
    raise Exception(f"AI Generation Error (All models exhausted): {last_err}")

# ---------------- MAIN RENDER FUNCTION ----------------
def render_exam_generator(api_key, default_logo_path, default_school_name="Saraswati Academy"):
    st.subheader("1. School & Examination Details")
    col_sc1, col_sc2 = st.columns([2, 1])
    with col_sc1:
        school_name = st.text_input("School / Institution Name", value=default_school_name)

    col1, col2, col3 = st.columns(3)
    with col1:
        class_level = st.selectbox("Select Class", ALL_CLASSES, index=11)

    is_junior_class = class_level in ["Nursery", "LKG", "UKG", "Class 1", "Class 2", "Class 3", "Class 4", "Class 5", "Class 6", "Class 7", "Class 8"]

    if is_junior_class:
        available_subjects = PRIMARY_MIDDLE_SUBJECTS
    elif class_level in ["Class 9", "Class 10"]:
        available_subjects = SECONDARY_SUBJECTS
    else:
        available_subjects = SENIOR_SECONDARY_SUBJECTS

    with col2:
        subject = st.selectbox("Select Subject", available_subjects)

    with col3:
        if is_junior_class:
            paper_standard = st.selectbox("Paper Standard", ["Standard School Level", "Easy / Activity Based", "Analytical / Olympiad"])
        else:
            paper_standard = st.selectbox(
                "Standard / Difficulty Level",
                ["PYQ (CBSE Board Previous Years Questions with Year Details)", "Standard CBSE (Medium)", "Basic / Easy", "Tough / Advanced"]
            )

    if is_junior_class:
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            total_target_marks = st.number_input("Target Total Marks", min_value=10, max_value=100, value=50, step=5)
        with col_m2:
            time_allowed = st.selectbox("Time Allowed", ["1 Hour", "1.5 Hours", "2 Hours", "2.5 Hours"], index=2)
    else:
        if "065" in subject or "083" in subject:
            total_target_marks = 70
            time_allowed = "3 Hours"
        elif "402" in subject or "417" in subject:
            total_target_marks = 50
            time_allowed = "2 Hours"
        elif "Physics" in subject or "Chemistry" in subject or "Biology" in subject:
            total_target_marks = 70
            time_allowed = "3 Hours"
        elif "Accountancy" in subject or "Business Studies" in subject or "Economics" in subject:
            total_target_marks = 80
            time_allowed = "3 Hours"
        else:
            total_target_marks = 80
            time_allowed = "3 Hours"

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            st.number_input("Maximum Marks (Locked as per CBSE SQP)", value=total_target_marks, disabled=True)
        with col_m2:
            st.text_input("Time Allowed (Locked as per CBSE SQP)", value=time_allowed, disabled=True)

    syllabus = st.text_area(
        "✍️ Enter or Paste Syllabus / Topics / Chapters",
        placeholder="Paste syllabus topics or chapters here...",
        height=100
    )

    st.markdown("---")
    st.subheader("2. Question Paper Blueprint")

    junior_blueprint = {}
    if is_junior_class:
        st.write("👉 **Select Question Types & Customize:**")
        for q_type in JUNIOR_Q_TYPES:
            if f"chk_{q_type}" not in st.session_state:
                st.session_state[f"chk_{q_type}"] = True
            if f"m_{q_type}" not in st.session_state:
                st.session_state[f"m_{q_type}"] = DEFAULT_JUNIOR_MARKS[q_type]
            if f"c_{q_type}" not in st.session_state:
                st.session_state[f"c_{q_type}"] = 0

        def auto_distribute_junior_marks(target_marks):
            selected_types = [t for t in JUNIOR_Q_TYPES if st.session_state.get(f"chk_{t}", False)]
            if not selected_types:
                return
            for t in JUNIOR_Q_TYPES:
                st.session_state[f"c_{t}"] = 0
            remaining = target_marks
            for t in selected_types:
                m = st.session_state.get(f"m_{t}", DEFAULT_JUNIOR_MARKS[t])
                if remaining >= m * 2:
                    st.session_state[f"c_{t}"] = 2
                    remaining -= (m * 2)
                elif remaining >= m:
                    st.session_state[f"c_{t}"] = 1
                    remaining -= m
            while remaining > 0:
                allocated = False
                for t in selected_types:
                    m = st.session_state.get(f"m_{t}", DEFAULT_JUNIOR_MARKS[t])
                    if remaining >= m:
                        st.session_state[f"c_{t}"] += 1
                        remaining -= m
                        allocated = True
                if not allocated:
                    break

        cols_h = st.columns([1, 4, 2, 2, 2])
        cols_h[0].write("**Include**")
        cols_h[1].write("**Question Type**")
        cols_h[2].write("**Marks / Q**")
        cols_h[3].write("**No. of Qs**")
        cols_h[4].write("**Total Marks**")

        current_calculated = 0
        for q_type in JUNIOR_Q_TYPES:
            col_chk, col_label, col_m, col_c, col_tot = st.columns([1, 4, 2, 2, 2])
            with col_chk:
                enabled = st.checkbox("", key=f"chk_{q_type}", label_visibility="collapsed")
            with col_label:
                st.write(f"**{q_type}**" if enabled else f"~{q_type}~")
            with col_m:
                marks = st.selectbox(f"marks_sel_{q_type}", [1, 2, 3, 4, 5, 6], key=f"m_{q_type}", disabled=not enabled, label_visibility="collapsed")
            with col_c:
                count = st.number_input(f"count_in_{q_type}", min_value=0, max_value=50, key=f"c_{q_type}", disabled=not enabled, label_visibility="collapsed")

            row_total = (marks * count) if enabled else 0
            current_calculated += row_total
            with col_tot:
                st.write(f"**{row_total}**")

            if enabled and count > 0:
                junior_blueprint[q_type] = {"marks": marks, "count": count}

        col_btn1, col_btn2 = st.columns([2, 2])
        with col_btn1:
            if current_calculated == total_target_marks:
                st.success(f"✅ Marks Matched: **{current_calculated} / {total_target_marks}**")
            else:
                st.warning(f"⚠️ Current Total: **{current_calculated}** | Target: **{total_target_marks}**")
        with col_btn2:
            st.button("🪄 Auto-Distribute Marks", on_click=auto_distribute_junior_marks, args=(total_target_marks,))
    else:
        if "065" in subject or "083" in subject:
            st.info("📋 **Official CBSE SQP Pattern (37 Questions):** Sec A (21 Qs × 1M), Sec B (7 Qs × 2M), Sec C (4 Qs × 3M), Sec D (2 Case Studies × 4M), Sec E (3 Qs × 5M).")
        elif "402" in subject or "417" in subject:
            st.info("📋 **Official CBSE Skill Blueprint (21 Questions):** Section A (Objective - 24 Marks) & Section B (Subjective - 26 Marks).")
        else:
            st.info("📋 **Official CBSE SQP Structure:** Sections A to E with MCQs, Assertion-Reason, VSA, SA, LA & Case Studies.")

    st.markdown("---")

    if st.button("🚀 Generate Examination Paper & Export DOCX", type="primary", use_container_width=True):
        if not api_key:
            st.error("❌ Application License key not configured. Please contact administrator.")
        elif not syllabus.strip():
            st.error("❌ Please enter syllabus or topics.")
        elif is_junior_class and not junior_blueprint:
            st.error("❌ Select at least one question type.")
        elif is_junior_class and current_calculated != total_target_marks:
            st.error(f"❌ Junior Blueprint Total ({current_calculated}) must match Target ({total_target_marks}).")
        else:
            try:
                with st.spinner(f"🧠 Generating authentic CBSE 3-Column Paper for {class_level} - {subject}..."):
                    system_prompt = (
                        "You are a Senior CBSE Examination Paper Setter adhering strictly to CBSE SQP blueprints. "
                        "CRITICAL: Do NOT write any thinking process or scratchpad. "
                        "Output ONLY the question paper text. Format every question starting with Q1., Q2., etc., followed by question text, options on separate lines, and marks at the end like [1] or [2]."
                    )

                    if is_junior_class:
                        bp_text = "\n".join([f"- Section '{k}': Exactly {v['count']} questions ({v['marks']} Marks each)" for k, v in junior_blueprint.items()])
                        user_prompt = f"Class: {class_level} | Subject: {subject} | Syllabus: {syllabus} | Total Marks: {total_target_marks} | Time: {time_allowed}\nBlueprint:\n{bp_text}"
                    else:
                        is_pyq = "PYQ" in paper_standard
                        pyq_tag = "Include authentic CBSE PYQ tags like [CBSE 2024] at the end of questions." if is_pyq else ""
                        if "402" in subject or "417" in subject:
                            user_prompt = f"Official CBSE Skill Paper: Class: {class_level}, Subject: {subject}, Syllabus: {syllabus}, Total Marks: 50, Time: 2 Hours. {pyq_tag}. Section A Objective (24 Marks) Q1 to Q5. Section B Subjective (26 Marks) Q6 to Q21."
                        elif "065" in subject or "083" in subject:
                            user_prompt = f"Official CBSE Paper: Class: {class_level}, Subject: {subject}, Syllabus: {syllabus}, Total Marks: 70, Time: 3 Hours. {pyq_tag}. 37 Questions (Sec A to E)."
                        else:
                            user_prompt = f"Official CBSE Paper: Class: {class_level}, Subject: {subject}, Syllabus: {syllabus}, Total Marks: {total_target_marks}, Time: {time_allowed}. {pyq_tag}. Sections A to E format."

                    generated_paper_text = run_gemini_paper_generator(api_key, system_prompt, user_prompt)

                filepath, filename = export_cbse_sqp_docx(
                    school_name=school_name,
                    class_level=class_level,
                    subject=subject,
                    total_marks=total_target_marks,
                    time_allowed=time_allowed,
                    raw_content=generated_paper_text,
                    logo_path=default_logo_path
                )

                st.success(f"🎉 Official CBSE 3-Column Paper generated successfully!")

                with open(filepath, "rb") as f:
                    st.download_button(
                        label="📥 Download Official Word Document (.docx)",
                        data=f,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

                st.markdown("---")
                if default_logo_path and os.path.exists(default_logo_path):
                    st.image(default_logo_path, width=100)

                st.markdown(
                    f"""
                    <div style="border: 2px solid #333; padding: 25px; border-radius: 6px; background-color: #ffffff; color: #111; font-family: 'Times New Roman', serif;">
                        <div style="text-align: center;">
                            <h2 style="margin: 0; font-size: 22px; font-weight: bold;">{school_name.upper()}</h2>
                            <h4 style="margin: 5px 0; font-size: 16px;">SAMPLE QUESTION PAPER (SESSION 2026-2027)</h4>
                            <h3 style="margin: 5px 0; color: #1a365d; font-size: 18px;">{class_level.upper()} — {subject.upper()}</h3>
                            <p style="margin: 5px 0; font-size: 14px;"><b>Time Allowed:</b> {time_allowed.upper()} &nbsp;&nbsp;|&nbsp;&nbsp; <b>Maximum Marks:</b> {total_target_marks}</p>
                        </div>
                        <hr style="margin: 15px 0;">
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                st.markdown(generated_paper_text)

            except Exception as err:
                st.error(f"Error: {str(err)}")
